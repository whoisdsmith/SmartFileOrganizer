import os
import io
import datetime
import logging
from typing import Dict, Optional, Tuple, Any, List

# Configure logging
logger = logging.getLogger(__name__)

# Handle imports with graceful fallbacks
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    logger.warning("pandas not available - CSV and Excel parsing will be limited")
    PANDAS_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    logger.warning("BeautifulSoup not available - HTML parsing will be limited")
    BS4_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - DOCX parsing will be limited")
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    logger.warning("chardet not available - character encoding detection will be limited")
    CHARDET_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 not available - PDF parsing will be limited")
    PYPDF2_AVAILABLE = False

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    logger.warning("Pillow not available - image processing will be limited")
    PIL_AVAILABLE = False

# Import media handling libraries
try:
    from pydub import AudioSegment
    import ffmpeg
    MEDIA_SUPPORT = True
except ImportError:
    logger.warning("Media libraries not available - audio/video processing will be limited")
    MEDIA_SUPPORT = False

# Import OCR service
from .ocr_service import OCRService


class FileParser:
    """
    Class for parsing different file types and extracting text content
    """

    def __init__(self, config: Optional[Dict] = None):
        """Initialize FileParser with optional configuration."""
        self.config = config or {}
        self.ocr_service = OCRService(self.config.get('ocr_config', {}))

    def extract_text(self, file_path, file_ext):
        """
        Extract text content from various file types

        Args:
            file_path: Path to the file
            file_ext: File extension (including the dot)

        Returns:
            Extracted text content as a string
        """
        if file_ext == '.csv':
            return self._parse_csv(file_path)
        elif file_ext == '.xlsx':
            return self._parse_excel(file_path)
        elif file_ext == '.html':
            return self._parse_html(file_path)
        elif file_ext == '.md':
            return self._parse_markdown(file_path)
        elif file_ext == '.txt':
            return self._parse_text(file_path)
        elif file_ext == '.docx':
            return self._parse_docx(file_path)
        elif file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            return self._parse_image(file_path)
        elif file_ext.lower() in ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a']:
            return self._parse_audio(file_path)
        elif file_ext.lower() in ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.webm', '.flv']:
            return self._parse_video(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")

    def extract_metadata(self, file_path, file_ext):
        """
        Extract metadata from files

        Args:
            file_path: Path to the file
            file_ext: File extension (including the dot)

        Returns:
            Dictionary containing metadata
        """
        # Basic file metadata
        file_stat = os.stat(file_path)
        metadata = {
            'file_name': os.path.basename(file_path),
            'file_path': file_path,
            'file_size': file_stat.st_size,
            'creation_time': datetime.datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            'modification_time': datetime.datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            'file_extension': file_ext,
        }

        # Extract file type specific metadata
        try:
            if file_ext == '.pdf':
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif file_ext == '.docx':
                docx_metadata = self._extract_docx_metadata(file_path)
                metadata.update(docx_metadata)
            elif file_ext.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
                image_metadata = self._extract_image_metadata(file_path)
                metadata.update(image_metadata)
            elif file_ext.lower() in ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a']:
                if MEDIA_SUPPORT:
                    audio_metadata = self._extract_audio_metadata(file_path)
                    metadata.update(audio_metadata)
            elif file_ext.lower() in ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.webm', '.flv']:
                if MEDIA_SUPPORT:
                    video_metadata = self._extract_video_metadata(file_path)
                    metadata.update(video_metadata)
        except Exception as e:
            metadata['metadata_error'] = str(e)

        return metadata

    def _parse_csv(self, file_path):
        """Parse CSV file content"""
        try:
            # Check if required libraries are available
            if not PANDAS_AVAILABLE:
                # Fallback to simple reading if pandas is not available
                logger.info("Using fallback for CSV parsing because pandas is not available")
                return self._parse_text(file_path)
            
            # Try to detect encoding
            encoding = 'utf-8'  # Default encoding
            if CHARDET_AVAILABLE:
                with open(file_path, 'rb') as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result['encoding']
            
            # Read CSV file with pandas
            df = pd.read_csv(file_path, encoding=encoding)

            # Convert to string representation
            text = df.to_string(index=False, max_rows=100)

            # Indicate if truncated
            if len(df) > 100:
                text += f"\n\n[File contains {len(df)} rows, showing first 100 only]"

            return text
        except Exception as e:
            logger.warning(f"Error parsing CSV file: {str(e)}")
            # Fallback to simple reading
            return self._parse_text(file_path)

    def _parse_excel(self, file_path):
        """Parse Excel file content"""
        try:
            # Check if pandas is available
            if not PANDAS_AVAILABLE:
                logger.info("Using fallback for Excel parsing because pandas is not available")
                return "[Excel parsing requires pandas library which is not available]"
            
            # Get sheet names
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            text = ""

            # Read each sheet
            for sheet in sheet_names[:3]:  # Limit to first 3 sheets
                df = pd.read_excel(file_path, sheet_name=sheet)

                text += f"Sheet: {sheet}\n"
                text += df.to_string(index=False, max_rows=50)

                # Indicate if truncated
                if len(df) > 50:
                    text += f"\n[Sheet contains {len(df)} rows, showing first 50 only]\n"

                text += "\n\n"

            # Indicate if more sheets exist
            if len(sheet_names) > 3:
                text += f"[File contains {len(sheet_names)} sheets, showing first 3 only]"

            return text
        except Exception as e:
            logger.warning(f"Error parsing Excel file: {str(e)}")
            return f"Error parsing Excel file: {str(e)}"

    def _parse_html(self, file_path):
        """Parse HTML file content"""
        try:
            # Check if BeautifulSoup is available
            if not BS4_AVAILABLE:
                logger.info("Using fallback for HTML parsing because BeautifulSoup is not available")
                return self._parse_text(file_path)
            
            # Detect encoding
            encoding = 'utf-8'  # Default encoding
            if CHARDET_AVAILABLE:
                with open(file_path, 'rb') as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result['encoding']
            
            # Read HTML file
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                html_content = f.read()

            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()

            # Get text
            text = soup.get_text(separator='\n')

            # Clean up multiple line breaks
            import re
            text = re.sub(r'\n+', '\n\n', text).strip()

            return text
        except Exception as e:
            logger.warning(f"Error parsing HTML file: {str(e)}")
            # Fallback to simple reading
            return self._parse_text(file_path)

    def _parse_markdown(self, file_path):
        """Parse Markdown file content"""
        return self._parse_text(file_path)

    def _parse_text(self, file_path):
        """Parse plain text file content"""
        try:
            # Detect encoding
            encoding = 'utf-8'  # Default encoding
            if CHARDET_AVAILABLE:
                with open(file_path, 'rb') as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result['encoding'] or 'utf-8'

            # Read text file
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.warning(f"Error parsing text file: {str(e)}")
            return f"Error parsing text file: {str(e)}"

    def _parse_docx(self, file_path):
        """Parse Word document content"""
        try:
            # Check if python-docx is available
            if not DOCX_AVAILABLE:
                logger.info("Using fallback for DOCX parsing because python-docx is not available")
                return "[DOCX parsing requires python-docx library which is not available]"
            
            # Open the document
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(' | '.join(row_text))

            return '\n'.join(full_text)
        except Exception as e:
            logger.warning(f"Error parsing Word document: {str(e)}")
            return f"Error parsing Word document: {str(e)}"

    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Parse PDF file and extract text content.
        For image-based PDFs, OCR is used to extract text.

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (extracted text, metadata)
        """
        text = ""
        metadata = {}

        # Check if PyPDF2 is available
        if not PYPDF2_AVAILABLE:
            logger.info("Using fallback for PDF parsing because PyPDF2 is not available")
            return "[PDF parsing requires PyPDF2 library which is not available]", {"error": "PyPDF2 not available"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted,
                    'version': pdf_reader.pdf_version
                }

                # Try to extract text directly first
                for page in pdf_reader.pages:
                    page_text = page.extract_text()

                    # If page has no extractable text, it might be image-based
                    if not page_text.strip():
                        # Get the OCR results for this page
                        ocr_results = self.ocr_service.process_pdf(file_path)

                        # Combine OCR results
                        for result in ocr_results:
                            if result['confidence'] > self.config.get('ocr_confidence_threshold', 50):
                                text += result['text'] + "\n"

                        # Add OCR metadata
                        metadata['ocr_used'] = True
                        if ocr_results:  # Check if there are any OCR results
                            metadata['ocr_confidence'] = sum(
                                r['confidence'] for r in ocr_results) / len(ocr_results)
                            metadata['ocr_languages'] = list(
                                set(r['language'] for r in ocr_results))
                    else:
                        text += page_text + "\n"

        except Exception as e:
            logger.warning(f"Error parsing PDF {file_path}: {str(e)}")
            return "", {}

        return text.strip(), metadata
        
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing PDF metadata
        """
        if not PYPDF2_AVAILABLE:
            return {"error": "PDF metadata extraction requires PyPDF2 library."}
            
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Basic PDF info
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted,
                    'pdf_version': pdf_reader.pdf_version
                }
                
                # Document information if available
                if pdf_reader.metadata:
                    for key, value in pdf_reader.metadata.items():
                        # Clean up the key name (remove leading /)
                        clean_key = key
                        if isinstance(key, str) and key.startswith('/'):
                            clean_key = key[1:]
                        metadata[f"pdf_{clean_key}"] = value
                
                return metadata
        except Exception as e:
            return {"error": f"Error extracting PDF metadata: {str(e)}"}
            
    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing DOCX metadata
        """
        if not DOCX_AVAILABLE:
            return {"error": "DOCX metadata extraction requires python-docx library."}
            
        try:
            doc = docx.Document(file_path)
            
            # Basic document properties
            properties = doc.core_properties
            
            metadata = {
                'docx_title': properties.title,
                'docx_author': properties.author,
                'docx_created': properties.created.isoformat() if properties.created else None,
                'docx_modified': properties.modified.isoformat() if properties.modified else None,
                'docx_last_modified_by': properties.last_modified_by,
                'docx_revision': properties.revision,
                'docx_category': properties.category,
                'docx_comments': properties.comments,
                'docx_keywords': properties.keywords,
                'docx_subject': properties.subject,
                'docx_paragraph_count': len(doc.paragraphs),
                'docx_table_count': len(doc.tables)
            }
            
            # Remove None values
            return {k: v for k, v in metadata.items() if v is not None}
        except Exception as e:
            return {"error": f"Error extracting DOCX metadata: {str(e)}"}

    def _parse_image(self, file_path):
        """
        Extract any text content from image (placeholder for OCR integration)

        Args:
            file_path: Path to the image file

        Returns:
            String with basic image information (no OCR yet)
        """
        # Check if PIL is available
        if not PIL_AVAILABLE:
            logger.info("Using fallback for image parsing because PIL is not available")
            return "[Image parsing requires PIL/Pillow library which is not available]"
            
        try:
            with Image.open(file_path) as img:
                # Basic image information
                width, height = img.size
                format_name = img.format
                mode = img.mode

                # Create a basic text representation
                text = f"Image: {os.path.basename(file_path)}\n"
                text += f"Dimensions: {width}x{height} pixels\n"
                text += f"Format: {format_name}\n"
                text += f"Color Mode: {mode}\n"

                # Add EXIF data summary if available
                exif_data = self._extract_image_metadata(file_path)
                if exif_data:
                    text += "\nImage Metadata:\n"
                    for key, value in exif_data.items():
                        if key not in ['filename', 'file_size', 'created_time', 'modified_time', 'file_extension']:
                            text += f"{key}: {value}\n"

                return text
        except Exception as e:
            logger.warning(f"Error parsing image: {str(e)}")
            return f"Error parsing image: {str(e)}"

    def _parse_audio(self, file_path):
        """
        Parse audio file and return a text representation

        Args:
            file_path: Path to the audio file

        Returns:
            Text representation of the audio file (metadata as text)
        """
        if not MEDIA_SUPPORT:
            return "Audio file parsing requires pydub and ffmpeg libraries."

        try:
            # Extract metadata
            metadata = self._extract_audio_metadata(file_path)

            # Convert metadata to text representation
            text_parts = [
                f"Audio File: {os.path.basename(file_path)}",
                f"Duration: {metadata.get('duration_seconds', 0):.2f} seconds",
                f"Channels: {metadata.get('channels', 'Unknown')}",
                f"Sample Rate: {metadata.get('sample_rate_hz', 'Unknown')} Hz",
                f"Bit Rate: {metadata.get('bitrate', 'Unknown')} bps",
            ]

            # Add ID3 tags if available
            for tag_name, tag_value in metadata.items():
                if tag_name.startswith('id3_'):
                    text_parts.append(
                        f"{tag_name.replace('id3_', '')}: {tag_value}")

            return "\n".join(text_parts)

        except Exception as e:
            return f"Error parsing audio file: {str(e)}"

    def _parse_video(self, file_path):
        """
        Parse video file and return a text representation

        Args:
            file_path: Path to the video file

        Returns:
            Text representation of the video file (metadata as text)
        """
        if not MEDIA_SUPPORT:
            return "Video file parsing requires ffmpeg library."

        try:
            # Extract metadata
            metadata = self._extract_video_metadata(file_path)

            # Convert metadata to text representation
            text_parts = [
                f"Video File: {os.path.basename(file_path)}",
                f"Duration: {metadata.get('duration_seconds', 0):.2f} seconds",
                f"Resolution: {metadata.get('width', 'Unknown')}x{metadata.get('height', 'Unknown')}",
                f"Video Codec: {metadata.get('video_codec', 'Unknown')}",
                f"Audio Codec: {metadata.get('audio_codec', 'Unknown')}",
                f"Frame Rate: {metadata.get('frame_rate', 'Unknown')} fps",
                f"Bit Rate: {metadata.get('bitrate', 'Unknown')} bps",
            ]

            return "\n".join(text_parts)

        except Exception as e:
            return f"Error parsing video file: {str(e)}"

    def _extract_audio_metadata(self, file_path):
        """
        Extract metadata from an audio file

        Args:
            file_path: Path to the audio file

        Returns:
            Dictionary containing audio metadata
        """
        if not MEDIA_SUPPORT:
            return {"error": "Audio metadata extraction requires pydub library."}

        try:
            # Load audio file
            audio = AudioSegment.from_file(file_path)

            # Extract basic metadata
            metadata = {
                'duration_seconds': len(audio) / 1000,
                'channels': audio.channels,
                'sample_width_bytes': audio.sample_width,
                'sample_rate_hz': audio.frame_rate,
                'frame_width': audio.frame_width,
                'bitrate': int((audio.frame_rate * audio.frame_width * audio.channels * 8)),
                'file_size_bytes': os.path.getsize(file_path),
            }

            # TODO: Extract ID3 tags for MP3 files
            # This would require a library like mutagen

            return metadata

        except Exception as e:
            return {"error": str(e)}

    def _extract_video_metadata(self, file_path):
        """
        Extract metadata from a video file

        Args:
            file_path: Path to the video file

        Returns:
            Dictionary containing video metadata
        """
        if not MEDIA_SUPPORT:
            return {"error": "Video metadata extraction requires ffmpeg library."}

        try:
            # Use ffmpeg to get video metadata
            probe = ffmpeg.probe(file_path)

            # Extract video stream info
            video_info = next((stream for stream in probe['streams']
                              if stream['codec_type'] == 'video'), None)

            # Extract audio stream info
            audio_info = next((stream for stream in probe['streams']
                              if stream['codec_type'] == 'audio'), None)

            # Extract format info
            format_info = probe['format']

            # Build metadata dictionary
            metadata = {
                'duration_seconds': float(format_info.get('duration', 0)),
                'file_size_bytes': os.path.getsize(file_path),
                'format_name': format_info.get('format_name', ''),
                'bitrate': int(format_info.get('bit_rate', 0)),
            }

            # Add video stream info if available
            if video_info:
                metadata.update({
                    'width': int(video_info.get('width', 0)),
                    'height': int(video_info.get('height', 0)),
                    'video_codec': video_info.get('codec_name', ''),
                    'video_codec_long': video_info.get('codec_long_name', ''),
                    'frame_rate': self._calculate_frame_rate(video_info),
                    'aspect_ratio': video_info.get('display_aspect_ratio', ''),
                })

            # Add audio stream info if available
            if audio_info:
                metadata.update({
                    'audio_codec': audio_info.get('codec_name', ''),
                    'audio_codec_long': audio_info.get('codec_long_name', ''),
                    'audio_channels': int(audio_info.get('channels', 0)),
                    'audio_sample_rate': int(audio_info.get('sample_rate', 0)),
                })

            return metadata

        except Exception as e:
            return {"error": str(e)}

    def _calculate_frame_rate(self, video_info):
        """
        Calculate the frame rate from ffmpeg video stream info

        Args:
            video_info: Dictionary containing video stream information

        Returns:
            Frame rate as a float
        """
        # Try to get frame rate from avg_frame_rate
        if 'avg_frame_rate' in video_info:
            try:
                num, den = video_info['avg_frame_rate'].split('/')
                return float(num) / float(den)
            except (ValueError, ZeroDivisionError):
                pass

        # Try to get frame rate from r_frame_rate
        if 'r_frame_rate' in video_info:
            try:
                num, den = video_info['r_frame_rate'].split('/')
                return float(num) / float(den)
            except (ValueError, ZeroDivisionError):
                pass

        return 0.0

    def _extract_image_metadata(self, file_path):
        """
        Extract metadata from image files including EXIF data

        Args:
            file_path: Path to the image file

        Returns:
            Dictionary with image metadata
        """
        metadata = {}
        
        # Check if PIL is available
        if not PIL_AVAILABLE:
            return {"error": "Image metadata extraction requires PIL/Pillow library."}

        try:
            with Image.open(file_path) as img:
                # Basic image properties
                metadata['image_width'], metadata['image_height'] = img.size
                metadata['image_format'] = img.format
                metadata['image_mode'] = img.mode

                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        # Process standard EXIF tags
                        for tag_id, value in exif.items():
                            tag = TAGS.get(tag_id, tag_id)

                            # Handle special cases
                            if tag == 'GPSInfo':
                                gps_data = {}
                                for gps_tag_id, gps_value in value.items():
                                    gps_tag = GPSTAGS.get(
                                        gps_tag_id, gps_tag_id)
                                    gps_data[gps_tag] = gps_value

                                # Calculate latitude and longitude if available
                                if 'GPSLatitude' in gps_data and 'GPSLatitudeRef' in gps_data:
                                    lat = self._convert_to_degrees(
                                        gps_data['GPSLatitude'])
                                    if gps_data['GPSLatitudeRef'] == 'S':
                                        lat = -lat
                                    metadata['gps_latitude'] = lat

                                if 'GPSLongitude' in gps_data and 'GPSLongitudeRef' in gps_data:
                                    lon = self._convert_to_degrees(
                                        gps_data['GPSLongitude'])
                                    if gps_data['GPSLongitudeRef'] == 'W':
                                        lon = -lon
                                    metadata['gps_longitude'] = lon

                                if 'GPSAltitude' in gps_data:
                                    metadata['gps_altitude'] = float(
                                        gps_data['GPSAltitude'])

                                metadata['gps_data'] = gps_data
                            elif tag == 'DateTime':
                                metadata['date_time'] = value
                            elif tag == 'DateTimeOriginal':
                                metadata['date_time_original'] = value
                            elif tag == 'DateTimeDigitized':
                                metadata['date_time_digitized'] = value
                            elif tag == 'Make':
                                metadata['camera_make'] = value
                            elif tag == 'Model':
                                metadata['camera_model'] = value
                            elif tag == 'XResolution':
                                metadata['x_resolution'] = float(value)
                            elif tag == 'YResolution':
                                metadata['y_resolution'] = float(value)
                            elif tag == 'ExposureTime':
                                metadata['exposure_time'] = str(value)
                            elif tag == 'FNumber':
                                metadata['f_number'] = float(value)
                            elif tag == 'ISOSpeedRatings':
                                metadata['iso_speed'] = value
                            elif tag == 'FocalLength':
                                metadata['focal_length'] = float(value)
                            else:
                                # Store other tags with proper formatting
                                if isinstance(value, bytes):
                                    try:
                                        value = value.decode('utf-8')
                                    except:
                                        value = str(value)
                                metadata[f'exif_{tag.lower()}'] = value
        except Exception as e:
            logger.warning(f"Error extracting image metadata: {str(e)}")
            metadata['exif_error'] = str(e)

        return metadata

    def _convert_to_degrees(self, value):
        """
        Helper method to convert GPS coordinates from EXIF format to decimal degrees

        Args:
            value: EXIF GPS coordinate value (degrees, minutes, seconds)

        Returns:
            Decimal degrees
        """
        degrees = float(value[0])
        minutes = float(value[1])
        seconds = float(value[2])

        return degrees + (minutes / 60.0) + (seconds / 3600.0)
